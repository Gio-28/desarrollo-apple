"""
Rellena la plantilla Word (template.docx) con los datos de un ContratoTurismo.

La plantilla se preserva intacta (logos, tablas, las 21 clausulas legales, firmas):
solo se escribe texto en las celdas/parrafos que en el documento original estan en blanco,
mas algunas filas nuevas insertadas en la tabla de confirmacion de reserva (habitacion,
numero de personas, resumen de pagos) que no existian en la plantilla original.
Este modulo conoce la estructura EXACTA de esas celdas (fue inspeccionada a mano), por lo que
si algun dia se edita el template.docx en Word, estos indices deben revisarse.
"""

import base64
import copy
import datetime
from io import BytesIO
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.documents.contrato_turismo.schema import ContratoTurismo
from app.documents.contrato_turismo.text_parser import _normalize_itinerario_lines

TEMPLATE_PATH = Path(__file__).parent / "template.docx"

MESES_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]

BLACK = RGBColor(0, 0, 0)


def _force_black(run) -> None:
    run.font.color.rgb = BLACK


# --------------------------------------------------------------------------
# Helpers genericos de edicion de un .docx ya cargado
# --------------------------------------------------------------------------

def _set_paragraph_text(paragraph: Paragraph, text: str, bold: bool | None = None) -> None:
    """Escribe 'text' en el primer run del parrafo (limpiando el resto) y fuerza que el
    texto quede en negro, sin importar que color tuviera el run original en la plantilla
    (evita heredar colores de placeholders). Si 'bold' no es None, tambien fuerza la
    negrilla de ese run."""
    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        run = paragraph.add_run(text)
    _force_black(run)
    if bold is not None:
        run.bold = bold


def _set_cell_text(cell: _Cell, value: str, paragraph_index: int = 0, bold: bool | None = None) -> None:
    paragraphs = cell.paragraphs
    if paragraph_index >= len(paragraphs):
        return
    _set_paragraph_text(paragraphs[paragraph_index], value, bold=bold)


def _set_cell_label_value(cell: _Cell, label: str, value: str, paragraph_index: int = 0) -> None:
    """Escribe 'label' + 'value' respetando el formato run-por-run que ya trae la
    plantilla en estas celdas: el/los primeros runs son la etiqueta en negrilla y el
    ULTIMO run (reservado para el dato) no esta en negrilla. Si se sobreescribiera todo
    el parrafo con un solo run heredado de la etiqueta, el valor quedaria tambien en
    negrilla -- por eso etiqueta y valor se escriben en runs separados."""
    paragraphs = cell.paragraphs
    if paragraph_index >= len(paragraphs):
        return
    paragraph = paragraphs[paragraph_index]
    runs = paragraph.runs
    if len(runs) < 2:
        _set_paragraph_text(paragraph, f"{label}{value}")
        return
    runs[0].text = label
    _force_black(runs[0])
    for r in runs[1:-1]:
        r.text = ""
    runs[-1].text = value
    runs[-1].bold = False
    _force_black(runs[-1])


def _append_to_last_run(paragraph: Paragraph, extra_text: str) -> None:
    if paragraph.runs:
        run = paragraph.runs[-1]
        run.text = (run.text or "") + extra_text
    else:
        run = paragraph.add_run(extra_text)
    _force_black(run)


def _insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p_elm = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new_p_elm)
    new_paragraph = Paragraph(new_p_elm, paragraph._parent)
    _set_paragraph_text(new_paragraph, text)
    return new_paragraph


def _fill_multiline_cell(cell: _Cell, text: str) -> None:
    lines = [line.strip() for line in text.split("\n") if line.strip()] or [""]
    existing = cell.paragraphs
    last_p = existing[-1]
    while len(existing) < len(lines):
        last_p = _insert_paragraph_after(last_p, "")
        existing = cell.paragraphs
    for i, p in enumerate(existing):
        _set_paragraph_text(p, lines[i] if i < len(lines) else "")


def _find_paragraph(paragraphs: list[Paragraph], contains: str) -> Paragraph | None:
    for p in paragraphs:
        if contains in p.text:
            return p
    return None


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._p
    element.getparent().remove(element)


def _replace_everywhere(doc, old_text: str, new_text: str) -> None:
    """Reemplaza el texto de un parrafo completo (uniendo todos sus runs) que coincida
    exactamente, en cualquier parte del documento -- incluye cuadros de texto/formas, que
    python-docx no expone via doc.paragraphs. Compara el parrafo COMPLETO (no cada <w:t>
    por separado) porque Word suele partir un mismo texto en varios runs dentro de un
    cuadro de texto (p.ej. "T-XXX" guardado como runs "T-" y "XXX"), y comparar run por
    run nunca encuentra coincidencia en ese caso."""
    for p_elm in doc.element.body.iter(qn("w:p")):
        paragraph = Paragraph(p_elm, doc)
        if paragraph.text.strip() == old_text:
            _set_paragraph_text(paragraph, new_text)


def _clone_row_after(table: Table, template_row_index: int, after_row_index: int):
    """Clona la fila template_row_index y la inserta despues de after_row_index.
    Devuelve el nuevo elemento <w:tr>."""
    tbl = table._tbl
    trs = tbl.tr_lst
    new_tr = copy.deepcopy(trs[template_row_index])
    trs[after_row_index].addnext(new_tr)
    return new_tr


# --------------------------------------------------------------------------
# Bloques especificos del contrato
# --------------------------------------------------------------------------

def _fill_empresa_cliente_tables(tables: list[Table], data: ContratoTurismo) -> None:
    empresa_table, cliente_table = tables[0], tables[1]

    # EMPRESA -> fila 7 (ASESOR COMERCIAL) columna valor
    _set_cell_text(empresa_table.rows[7].cells[1], data.asesor_comercial)

    # CLIENTE -> filas 0..6
    _set_cell_text(cliente_table.rows[0].cells[1], data.cliente_nombre)
    _set_cell_text(cliente_table.rows[1].cells[1], data.cliente_cedula)
    _set_cell_text(cliente_table.rows[2].cells[1], data.cliente_direccion)
    _set_cell_text(cliente_table.rows[3].cells[1], data.cliente_telefono)
    _set_cell_text(cliente_table.rows[4].cells[1], data.cliente_correo)
    _set_cell_text(cliente_table.rows[5].cells[1], data.destino)
    _set_cell_text(cliente_table.rows[6].cells[1], data.confirmacion_reserva.upper())


def _fill_payment_table(tables: list[Table], data: ContratoTurismo) -> None:
    pay_table = tables[2]

    # fila 0: VALOR TOTAL (parrafo 2) / FECHA LIMITE DE PAGO (parrafo 2, celda combinada col1-2)
    # bold=False: son los DATOS (no la etiqueta "VALOR TOTAL" / "FECHA LIMITE..." de arriba,
    # que no se toca), y esa etiqueta esta en negrilla en la plantilla -- sin forzarlo aqui
    # el valor la heredaria.
    _set_cell_text(pay_table.rows[0].cells[0], f"$ {data.valor_total}", paragraph_index=1, bold=False)
    _set_cell_text(pay_table.rows[0].cells[1], data.fecha_limite_pago, paragraph_index=1, bold=False)

    # filas 2 y 3 son la plantilla de cada abono (fecha | $valor). Se clona segun # de pagos.
    tbl = pay_table._tbl
    trs = tbl.tr_lst
    template_tr = copy.deepcopy(trs[2])
    tbl.remove(trs[2])
    tbl.remove(trs[3])

    for pago in data.pagos:
        new_tr = copy.deepcopy(template_tr)
        tbl.append(new_tr)

    # ahora que las filas estan en el arbol, usamos la API normal de python-docx para escribir texto
    for i, pago in enumerate(data.pagos):
        row = pay_table.rows[2 + i]
        _set_cell_text(row.cells[0], pago.fecha)
        _set_cell_text(row.cells[1], f"$ {pago.valor}")


def _fill_beneficiarios_adicionales(doc, data: ContratoTurismo) -> None:
    """Llena los blancos de la clausula 16 (VIAJEROS Y CLIENTES O BENEFICIARIOS
    ADICIONALES) con TODOS los pasajeros de la reserva (titular + acompañantes), la
    misma lista que se usa en la tabla de confirmacion de reserva: si viaja una sola
    persona queda solo el titular, y si viajan varias, el titular y cada acompañante."""
    blanks = [p for p in doc.paragraphs if p.text and set(p.text.strip()) == {"_"}]
    if not blanks:
        return

    pasajeros = data.pasajeros_reserva
    for i, p in enumerate(blanks):
        if i < len(pasajeros):
            b = pasajeros[i]
            _set_paragraph_text(p, f"{b.nombre}          C.C. {b.documento}")

    last_p = blanks[-1]
    for i in range(len(blanks), len(pasajeros)):
        b = pasajeros[i]
        last_p = _insert_paragraph_after(last_p, f"{b.nombre}          C.C. {b.documento}")


def _fill_signature_block(doc, data: ContratoTurismo) -> None:
    cc_p = _find_paragraph(doc.paragraphs, "C.C 1152683639")
    if cc_p:
        _append_to_last_run(cc_p, data.cliente_cedula)

    # ancla en "Gerente y R. Legal" (unico en el documento) para no confundir con el
    # encabezado "EL CLIENTE O CONTRATANTE" que aparece mucho antes, en la tabla de datos
    el_cliente_p = _find_paragraph(doc.paragraphs, "Gerente y R. Legal")
    if el_cliente_p:
        for run in el_cliente_p.runs:
            if "EL CLIENTE" in run.text:
                run.text = run.text.replace("EL CLIENTE", data.cliente_nombre)

    fecha_p = _find_paragraph(doc.paragraphs, "Una vez leído en su integridad")
    if fecha_p and len(fecha_p.runs) >= 6:
        hoy = datetime.date.today()
        fecha_p.runs[1].text = f" {hoy.day}"
        fecha_p.runs[2].text = " de "
        fecha_p.runs[3].text = MESES_ES[hoy.month - 1]
        fecha_p.runs[4].text = " de "
        fecha_p.runs[5].text = str(hoy.year)


def _fill_reservation_table(tables: list[Table], data: ContratoTurismo) -> None:
    res_table = tables[4]

    # Nota: el codigo de reserva vive en un cuadro de texto sobre esta tabla (no es una
    # celda normal); se reemplaza a nivel de documento completo en fill_contract().

    row0 = res_table.rows[0]
    _append_to_last_run(row0.cells[0].paragraphs[0], data.programa.upper())
    _set_cell_text(row0.cells[1], data.fecha_reserva.upper(), paragraph_index=1, bold=False)

    # filas nuevas: HABITACION / N. DE PERSONAS, y resumen de PAGOS -- clonadas de filas
    # existentes de la misma tabla para heredar bordes/estilo, insertadas despues de
    # "RESERVADO A" (fila 1) y antes de "HOTEL / CHECK IN / CHECK OUT" (fila 2 original).
    # ojo con los indices: cada insercion corre las filas siguientes, por eso el
    # template_row_index de cada paso se recalcula segun el estado DESPUES del paso anterior.
    _clone_row_after(res_table, template_row_index=0, after_row_index=1)  # -> HABITACION en indice 2 (patron 2 celdas)
    _clone_row_after(res_table, template_row_index=4, after_row_index=2)  # -> PAGOS header en indice 3 (patron 1 celda)
    _clone_row_after(res_table, template_row_index=4, after_row_index=3)  # -> PAGOS valores en indice 4 (patron HOTEL, 3 celdas)

    # las filas originales HOTEL/CHECKIN/CHECKOUT y DATOS DE PASAJEROS ahora quedaron
    # corridas 3 posiciones (se insertaron 3 filas nuevas antes de ellas).
    # cells[1] de esta fila se clono del patron "FECHA" (2 parrafos: etiqueta + valor);
    # se usa el primer parrafo para el texto y se ELIMINA el segundo (si no, queda una
    # linea en blanco sobrante en la celda).
    row_habitacion = res_table.rows[2]
    _set_cell_label_value(row_habitacion.cells[0], "HABITACIÓN: ", data.habitacion.upper())
    _set_cell_label_value(row_habitacion.cells[1], "N° DE PERSONAS: ", data.cantidad_personas.upper())
    if len(row_habitacion.cells[1].paragraphs) > 1:
        _remove_paragraph(row_habitacion.cells[1].paragraphs[1])

    row_pagos_header = res_table.rows[3]
    _set_cell_text(row_pagos_header.cells[0], "PAGOS")

    # fila de PAGOS-valores: se clona del patron de HOTEL/CHECKIN/CHECKOUT (3 celdas reales,
    # no una sola combinada), asi cada valor queda en su propia casilla separada.
    row_pagos_valores = res_table.rows[4]
    _set_cell_label_value(row_pagos_valores.cells[0], "VALOR TOTAL: $ ", data.valor_total.upper())
    _set_cell_label_value(row_pagos_valores.cells[2], "VALOR ABONADO: $ ", data.valor_abonado.upper())
    _set_cell_label_value(row_pagos_valores.cells[3], "VALOR RESTANTE: $ ", data.valor_restante.upper())

    row_hotel = res_table.rows[5]
    _set_cell_label_value(row_hotel.cells[0], "HOTEL:    ", data.hotel.upper())
    _set_cell_label_value(row_hotel.cells[2], "CHECK IN: ", data.check_in.upper())
    _set_cell_label_value(row_hotel.cells[3], "CHECK OUT: ", data.check_out.upper())

    # fila 8 (DATOS DE LOS PASAJEROS: NOMBRE, en indice 7) es la plantilla de encabezado con
    # 2 celdas reales (NOMBRE | DOCUMENTO). Se elimina la fila original combinada de un solo
    # pasajero y se clona una fila real de 2 celdas por cada pasajero, en su lugar.
    tbl = res_table._tbl
    tbl.remove(tbl.tr_lst[8])

    last_index = 7
    for pax in data.pasajeros_reserva:
        _clone_row_after(res_table, template_row_index=7, after_row_index=last_index)
        last_index += 1
        new_row = res_table.rows[last_index]
        _set_cell_text(new_row.cells[0], pax.nombre.upper())
        _set_cell_text(new_row.cells[2], pax.documento.upper())


def _fill_incluye_no_incluye(tables: list[Table], data: ContratoTurismo) -> None:
    incluye_table, no_incluye_table = tables[5], tables[6]
    _fill_multiline_cell(incluye_table.rows[1].cells[0], data.incluye)
    _fill_multiline_cell(no_incluye_table.rows[1].cells[0], data.no_incluye)


def _decode_data_uri(data_uri: str) -> bytes | None:
    if "," in data_uri:
        data_uri = data_uri.split(",", 1)[1]
    try:
        return base64.b64decode(data_uri)
    except Exception:  # noqa: BLE001
        return None


def _fill_itinerario_y_tiquetes(doc, data: ContratoTurismo) -> None:
    """Agrega, al final del documento (despues de "NO INCLUYE"), el itinerario en texto
    y/o la captura de los tiquetes aereos, si se proporcionaron. Ninguno de los dos
    existe en la plantilla original: se insertan como parrafos nuevos."""
    anchor = doc.tables[6]._tbl  # tabla "EL PROGRAMA NO INCLUYE", ultimo bloque de la plantilla

    def _paragraph_after(after_element, center: bool = False):
        p = doc.add_paragraph()
        after_element.addnext(p._p)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    if data.itinerario.strip():
        # se normaliza aqui de nuevo (ademas de en el parser del texto pegado) para que
        # tambien quede bien organizado si el itinerario se pego o se edito a mano
        # directamente en el campo de edicion, sin pasar por el parser: colapsa espacios
        # en blanco irregulares y descarta un titulo "ITINERARIO" duplicado si la persona
        # copio el texto con su propio encabezado.
        itinerario_lines = _normalize_itinerario_lines(data.itinerario.split("\n"))

        title_p = _paragraph_after(anchor, center=True)
        title_run = title_p.add_run("ITINERARIO")
        title_run.bold = True
        _force_black(title_run)
        anchor = title_p._p
        for line in itinerario_lines:
            body_p = _paragraph_after(anchor)
            if line.strip():
                _force_black(body_p.add_run(line))
            anchor = body_p._p

    if data.tiquetes_imagen.strip():
        image_bytes = _decode_data_uri(data.tiquetes_imagen)
        if image_bytes:
            title_p = _paragraph_after(anchor, center=True)
            title_run = title_p.add_run("TIQUETES AÉREOS")
            title_run.bold = True
            _force_black(title_run)
            anchor = title_p._p

            img_p = _paragraph_after(anchor, center=True)
            img_p.add_run().add_picture(BytesIO(image_bytes), width=Inches(5.5))
            anchor = img_p._p


def _dedupe_pasajeros(pasajeros: list) -> list:
    """Quita pasajeros duplicados (mismo nombre + documento, sin importar mayus/minus ni
    espacios de mas) antes de imprimirlos -- p.ej. si el titular termino listado dos veces
    porque el formato pegado ya lo incluye automaticamente y ademas se agrego a mano en el
    formulario, evita que salga repetido tanto en la clausula 16 como en la tabla de
    confirmacion de reserva."""
    seen = set()
    out = []
    for p in pasajeros:
        key = (p.nombre.strip().lower(), p.documento.strip().lower())
        if key in seen:
            continue
        seen.add(key)
        out.append(p)
    return out


# --------------------------------------------------------------------------
# Punto de entrada
# --------------------------------------------------------------------------

def fill_contract(data: ContratoTurismo) -> bytes:
    doc = docx.Document(str(TEMPLATE_PATH))
    tables = doc.tables

    data.pasajeros_reserva = _dedupe_pasajeros(data.pasajeros_reserva)

    _fill_empresa_cliente_tables(tables, data)
    _fill_payment_table(tables, data)
    _fill_beneficiarios_adicionales(doc, data)
    _fill_signature_block(doc, data)
    _fill_reservation_table(tables, data)
    _fill_incluye_no_incluye(tables, data)
    _fill_itinerario_y_tiquetes(doc, data)
    _replace_everywhere(doc, "T-XXX", data.confirmacion_reserva.upper())

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
